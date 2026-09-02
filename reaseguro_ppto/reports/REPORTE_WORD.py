from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

def generar_reporte_word(
    resumen_ln,
    archivo_docx,
    carpeta_graficas=None
):
    """
    resumen_ln: DataFrame con métricas por LN

    archivo_docx: ruta destino

    carpeta_graficas:
        carpeta donde existen imágenes:
        LN04001.png
        LN04002.png
        etc
    """

    doc = Document()

    # ==================================================
    # PORTADA
    # ==================================================

    titulo = doc.add_heading(
        'Validación Forecast 2027',
        level=1
    )

    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run(
        f"Generado: {datetime.now():%d/%m/%Y %H:%M}"
    )

    doc.add_page_break()

    # ==================================================
    # RESUMEN EJECUTIVO
    # ==================================================

    doc.add_heading(
        'Resumen Ejecutivo',
        level=1
    )

    total_primas = resumen_ln["PRIMAS_"].sum()

    total_sin = resumen_ln["SINIESTROS_"].sum()

    total_com = resumen_ln["COMISIONES_"].sum()

    doc.add_paragraph(
        f"""
Primas Forecast: {total_primas:,.0f}

Siniestros Forecast: {total_sin:,.0f}

Comisiones Forecast: {total_com:,.0f}
"""
    )

    # ==================================================
    # DETALLE POR LN
    # ==================================================

    for _, row in resumen_ln.iterrows():

        ln = row["Linea_negocio"]

        doc.add_page_break()

        doc.add_heading(
            f"Validación FCST 2027 | {ln}",
            level=1
        )

        # KPIs

        tabla = doc.add_table(
            rows=4,
            cols=2
        )

        tabla.cell(0,0).text = "Primas"
        tabla.cell(0,1).text = f"{row['PRIMAS_']:,.0f}"

        tabla.cell(1,0).text = "Siniestros"
        tabla.cell(1,1).text = f"{row['SINIESTROS_']:,.0f}"

        tabla.cell(2,0).text = "Comisiones"
        tabla.cell(2,1).text = f"{row['COMISIONES_']:,.0f}"

        ratio = (
            row["SINIESTROS_"]
            /
            row["PRIMAS_"]
            if row["PRIMAS_"] != 0
            else 0
        )

        tabla.cell(3,0).text = "Siniestralidad"
        tabla.cell(3,1).text = f"{ratio:.2%}"

        # Gráfica

        if carpeta_graficas:

            grafica = os.path.join(
                carpeta_graficas,
                f"{ln}.png"
            )

            if os.path.exists(grafica):

                doc.add_picture(
                    grafica,
                    width=Inches(5.5)
                )

        # Comentario ejecutivo

        doc.add_heading(
            "Análisis",
            level=2
        )

        doc.add_paragraph(
            f"""
La línea {ln} presenta primas por
{row['PRIMAS_']:,.0f},
siniestros por
{row['SINIESTROS_']:,.0f}
y comisiones por
{row['COMISIONES_']:,.0f}.
"""
        )

        # Preguntas automáticas

        doc.add_heading(
            "Consultas",
            level=2
        )

        preguntas = [

            "¿La estacionalidad es consistente con el histórico?",

            "¿Existen meses con concentración relevante de primas?",

            "¿Los siniestros son congruentes con el crecimiento esperado?",

            "¿Las comisiones evolucionan en línea con las primas?",

            "¿Existe alguna desviación material contra presupuesto?"
        ]

        for pregunta in preguntas:

            doc.add_paragraph(
                pregunta,
                style="List Bullet"
            )

    # ==================================================
    # GUARDAR
    # ==================================================

    doc.save(
        archivo_docx
    )

    print(
        f"Reporte generado: {archivo_docx}"
    )